"""Build a full Supplementary Information draft (Supplement_full_draft.docx).
Formatting modeled on the author's example journal supplements. Reads verified
tables from outputs/tables/. Does NOT overwrite the user's Supplement.docx.
"""
import re as _re   # aliased: Table S1 below uses a DataFrame named `re`
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from adtopo.config import BASE_DIR as ROOT, CBCL_MEDIATION_OUTCOMES
TAB = ROOT / 'outputs' / 'tables'
WP  = ROOT / 'code/05_behavior/within_person/derived/results_within_person_cbcl.txt'
OUT = ROOT / 'manuscript' / 'Supplement_full_draft.docx'

FONT = 'Times New Roman'

_SUP = str.maketrans('0123456789-', '⁰¹²³⁴⁵⁶⁷⁸⁹⁻')
def _sci(x):
    m, e = f"{float(x):.1e}".split('e')
    return f"{m} × 10{str(int(e)).translate(_SUP)}"

def fmt_p(p):
    p = float(p)
    if p < 1e-4: return _sci(p)
    if p < 0.001: return "< .001"
    return f"{p:.3f}".lstrip('0').replace('0.', '.') if p < 1 else f"{p:.3f}"

def fmt_q(q):
    q = float(q)
    if q < 0.001: return "< .001"
    return f"{q:.3f}".lstrip('0')

def fmt_bf(bf10, bf01):
    bf10 = float(bf10)
    if bf10 >= 1:
        if bf10 >= 1000:
            m, e = f"{bf10:.1e}".split('e'); return f"BF₁₀ = {m} × 10{str(int(e)).translate(_SUP)}"
        return f"BF₁₀ = {bf10:.0f}" if bf10 >= 10 else f"BF₁₀ = {bf10:.1f}"
    return f"BF₀₁ = {float(bf01):.1f}"

doc = Document()
st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(11)

def para(text='', *, bold=False, italic=False, size=11, align='left', space_after=6, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT, 'center': WD_ALIGN_PARAGRAPH.CENTER}[align]
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text); r.bold = bold; r.italic = italic
        r.font.name = FONT; r.font.size = Pt(size)
    return p

def heading(text, size=12):
    para(text, bold=True, size=size, space_after=4)

def note(label, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(label); r.italic = True; r.font.name = FONT; r.font.size = Pt(10)
    r2 = p.add_run(' ' + text); r2.font.name = FONT; r2.font.size = Pt(10)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].add_run(h).bold = True
        for run in hdr[i].paragraphs[0].runs: run.font.size = Pt(9); run.font.name = FONT
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for run in cells[i].paragraphs[0].runs: run.font.size = Pt(9); run.font.name = FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ══════════════════════════════════════════════════════ TITLE + CONTENTS
p = para('Supplementary Information', bold=True, size=14, align='center', space_after=4)
para('Early threat exposure expands the somato-cognitive action network (SCAN) in the developing cortex',
     italic=True, size=11, align='center', space_after=14)

heading('Contents', size=12)
toc = [
    'Supplementary Methods',
    ' Model specification, nesting, and convergence',
    'Supplementary Results',
    ' Representativeness of the analytic sample',
    ' Individual dimensions of adversity and Bayes factors',
    ' Associations with psychopathology (CBCL)',
    'Supplementary Tables',
    ' Table S1. Robustness of the threat–SCAN association across seven variance structures',
    ' Table S2. Design and nesting structure of the analytic sample',
    ' Table S3. Comparison of included and excluded participants',
    ' Table S4. Associations of the ten adversity factors with SCAN size, with Bayes factors',
    ' Table S5. Joint model of all ten adversity factors predicting SCAN size',
    ' Table S6. SCAN mediation of threat and year-6 CBCL psychopathology',
    ' Table S7. Within-person SCAN change and CBCL psychopathology',
]
for line in toc:
    para(line, bold=line.startswith('Supplementary'), size=11, space_after=2)
doc.add_page_break()

# ══════════════════════════════════════════════════════ SUPP METHODS
heading('Supplementary Methods', size=13)
heading('Adversity exposure factors and composites', size=11)
para("Early-life adversity was characterized by ten data-driven factors derived from the early-environment "
     "measures collected in the ABCD Study [reference], each z-scored across the analytic sample. Following "
     "the dimensional model of adversity, we grouped the ten factors a priori into three composites, each "
     "computed as the mean of its constituent z-scored factors: threat (experiences of harm or its "
     "anticipation)—physical trauma, family aggression, family conflict, and family anger; deprivation "
     "(the absence of expected cognitive and social input)—socioeconomic and neighborhood disadvantage, "
     "low primary caregiver support, low secondary caregiver support, and caregiver supervision; and "
     "unpredictability (instability of the caregiving environment)—caregiver psychopathology and caregiver "
     "substance use or separation.")
para("Three factors—family anger and primary and secondary caregiver support—were reverse-coded before "
     "compositing so that higher scores on every factor and every composite denote greater adversity. Each "
     "factor correlated positively with the composite to which it was assigned (r = 0.41–0.92; Figure S1). "
     "The three composites were themselves moderately intercorrelated (threat–deprivation r = 0.75; "
     "threat–unpredictability r = 0.73; deprivation–unpredictability r = 0.60), which motivates the "
     "multivariate model in which the composites compete to predict SCAN size (main text); Table S5 reports "
     "the corresponding model that enters all ten individual factors simultaneously.")
heading('Model specification, nesting, and convergence', size=11)
para("All associations between early-life adversity and cortical network topography, encroachment, and "
     "functional connectivity reported in the main text were estimated with ordinary least squares (OLS), "
     "adjusting for age, sex, and in-scanner head motion (mean framewise displacement), with study site "
     "entered as a set of fixed-effect indicators. Non-independence of siblings was accommodated with "
     "family-cluster-robust (sandwich) standard errors rather than a family random intercept.")
para("We adopted this specification for three reasons, each documented in Table S2. First, all three "
     "scanner manufacturers and all 29 scanner serial numbers were perfectly nested within study site "
     "(no site used more than one manufacturer, and no serial number spanned more than one site), so "
     "fixed site indicators fully absorb between-scanner as well as between-site variance; adding a "
     "scanner random effect on top of site is therefore unidentified. Second, families never crossed "
     "sites (0 of 4,058), so family is structurally nested within site. Third, only 456 families (11.2%) "
     "contributed more than one imaged sibling, so a family random-intercept variance is estimated at "
     "essentially zero (its covariance is singular in every model that includes it). We therefore handle "
     "sibling non-independence with family-clustered robust standard errors rather than a boundary-valued "
     "variance component.")
para("To confirm that this fixed-effects specification does not drive the findings, we refit the primary "
     "threat→SCAN association—both the single-predictor (bivariate) and the joint three-composite "
     "(multivariate) models—under seven variance structures (Table S1): (i) OLS with fixed site and "
     "family-cluster-robust standard errors (the specification reported throughout); (ii) OLS with fixed "
     "site and classical standard errors; (iii) a random site intercept; (iv) fixed site with a random "
     "family intercept; (v) crossed random site and family intercepts; (vi) a random site with family "
     "nested within site; and (vii) fixed site with an added scanner-serial random effect and a random "
     "family intercept. The standardized association was invariant across all seven specifications "
     "(bivariate partial r = 0.151–0.161; multivariate β = 0.00231–0.00233), and every model "
     "reached optimizer convergence. Specifications that include a family random intercept return a "
     "singular random-effects covariance, reflecting the near-zero family variance described above; this "
     "is a property of the sibling structure rather than a failure of estimation, and the fixed-effect "
     "inference on which all reported associations rest is unaffected.")

# ══════════════════════════════════════════════════════ SUPP RESULTS
heading('Supplementary Results', size=13)

heading('Representativeness of the analytic sample', size=11)
para("Relative to the remainder of the ABCD baseline cohort (n = 7,345), the analytic sample "
     "(n = 4,525) was modestly more likely to be female (52.0% vs. 45.3%), White (61.9% vs. 48.2%), "
     "and from higher-income and higher-education households (Table S3). All associations were small "
     "(bias-corrected Cramér's V = 0.065–0.164; Bergsma, 2013). This selection pattern is "
     "consistent with the well-documented association between usable MRI quality and sociodemographic "
     "characteristics in developmental neuroimaging, and we note it as a bound on generalizability in "
     "the main-text limitations.")

heading('Individual dimensions of adversity and Bayes factors', size=11)
para("The main text reports associations between the three adversity composites and SCAN size. Table S4 "
     "provides the parallel results for each of the ten individual adversity factors, together with Bayes "
     "factors. Nine of the ten factors were independently associated with SCAN size at baseline (all FDR "
     "q < .05); the exception was caregiver supervision. Bayes factors, computed as BIC approximations "
     "from the same fixed-site models (BF₁₀ = exp[(BIC_null − BIC_full)/2]), corroborated "
     "this pattern: the four threat-type factors provided very strong evidence for an association "
     "(BF₁₀ ranging from 5.7×10¹⁰ to 5.3×10¹⁶), whereas caregiver "
     "supervision—the one factor that was non-significant by FDR—provided positive evidence for "
     "the null (BF₀₁ = 3.0).")
para("Because the ten factors are correlated, we also fit a single model entering all ten simultaneously "
     "(Table S5). In this joint model, only family aggression remained independently associated with SCAN "
     "size after FDR correction (β = 0.0012, q < .001), consistent with substantial shared variance "
     "among the threat-type factors and supporting our decision to summarize them with the a priori "
     "threat, deprivation, and unpredictability composites rather than to interpret each factor in "
     "isolation.")

heading('Associations with psychopathology (CBCL)', size=11)
para("To complement the cognitive analyses, we tested whether baseline SCAN size mediated the association "
     "between threat and each of 14 Child Behavior Checklist (CBCL) syndrome and DSM-oriented subscales "
     "at year 6, using the same bootstrapped (5,000 family-clustered resamples), covariate-matched "
     "mediation framework applied to cognition (Table S6). No subscale showed a significant indirect "
     "effect: none survived FDR correction and every bootstrap confidence interval included zero "
     "(all bootstrap p > .05; all FDR q > .47). Consistent with this, within-person change in SCAN "
     "size did not track within-person change in any of the 14 subscales after FDR correction "
     "(Table S7): three subscales showed a nominal association (p < .05) but none survived correction "
     "(all q > .06). Across both cross-sectional and within-person tests, and after correction for "
     "multiple comparisons, we found no evidence that SCAN size relates to concurrent or developing "
     "psychopathology, underscoring the specificity of the SCAN–cognition association reported in the "
     "main text.")

doc.add_page_break()
heading('Supplementary Tables', size=13)

# ---- Table S1: RE-invariance
re = pd.read_csv(TAB / 'supp_table_re_invariance.csv')
para('Table S1. Robustness of the threat–SCAN association across seven variance structures',
     bold=True, size=11, space_after=4)
rowsS1 = []
for _, r in re.iterrows():
    est = f"{r['estimate']:.3f}" if r['estimate_type'] == 'partial_r' else f"{r['estimate']:.5f}"
    model = 'Bivariate' if 'Bivariate' in r['Model'] else 'Multivariate'
    etype = 'partial r' if r['estimate_type'] == 'partial_r' else 'β'
    rowsS1.append([model, r['Specification'], f"{etype} = {est}", f"{r['se']:.5f}",
                   fmt_p(r['p']), 'Yes' if r['converged'] else 'No'])
add_table(['Model', 'Specification', 'Estimate', 'SE', 'p', 'Converged'], rowsS1)
note('Note.', "Bivariate = single-predictor threat→SCAN model; Multivariate = threat estimate from "
     "the joint three-composite model. Specification (i), OLS with fixed site and family-cluster-robust "
     "standard errors, is the model reported throughout the main text. The standardized association is "
     "invariant across all seven specifications and every model converged.")

# ---- Table S2: design/nesting
dn = pd.read_csv(TAB / 'supp_design_nesting.csv')
para('Table S2. Design and nesting structure of the analytic sample', bold=True, size=11, space_after=4)
labelmap = {
 'n_study_sites': 'Study-site indicator levels',
 'n_scanner_manufacturers': 'Scanner manufacturers',
 'sites_with_>1_manufacturer': 'Sites using >1 manufacturer',
 'n_scanner_serials': 'Scanner serial numbers',
 'serials_spanning_>1_site': 'Serial numbers spanning >1 site',
 'max_serials_per_site': 'Maximum serial numbers per site',
 'n_families': 'Distinct families',
 'families_spanning_>1_site': 'Families spanning >1 site',
 'families_with_>=2_imaged_sibs': 'Families with ≥2 imaged siblings',
 'pct_participants_with_a_sib_in_sample': 'Participants with a sibling in sample (%)',
 'phase3_baseline_family_var~0_models': 'Models with singular family variance',
}
rowsS2 = [[labelmap.get(r['quantity'], r['quantity']), str(r['value']), r['interpretation']]
          for _, r in dn.iterrows()]
add_table(['Quantity', 'Value', 'Interpretation'], rowsS2)
note('Note.', "The study-site variable comprises 22 indicator levels—the 21 ABCD research sites plus "
     "one administrative level (n = 15) for participants not affiliated with a standard collection "
     "site—all entered as fixed-effect indicators. Because scanner manufacturer and serial number are "
     "perfectly nested within site, and families never cross sites, fixed site indicators absorb the "
     "scanner and between-site variance; the near-zero family variance motivates family-cluster-robust "
     "standard errors in place of a family random intercept.")

# ---- Table S3: included vs excluded
para('Table S3. Comparison of included and excluded participants', bold=True, size=11, space_after=4)
rowsS3 = [
 ['Sex (% female)', '45.3', '52.0', 'χ²(1) = 50.9', '9.6 × 10⁻¹³', '0.065'],
 ['Race/ethnicity (% White)', '48.2', '61.9', 'χ²(4) = 258.2', '1.1 × 10⁻⁵⁴', '0.153'],
 ['Household income (% ≥ $100k)', '41.6', '54.8', 'χ²(9) = 294.3', '4.3 × 10⁻⁵⁸', '0.164'],
 ['Parental education (% ≥ bachelor’s)', '47.2', '58.5', 'χ²(20) = 211.6', '5.5 × 10⁻³⁴', '0.133'],
]
add_table(['Characteristic', 'Excluded %', 'Included %', 'χ² (df)', 'p', "Cramér's V"], rowsS3)
note('Note.', "Excluded = remainder of the ABCD baseline cohort (n = 7,345); Included = analytic sample "
     "(n = 4,525). Income ≥ $100k combines ABCD categories 9–10; parental education ≥ "
     "bachelor’s combines categories 18–21. Cramér's V is bias-corrected (Bergsma, 2013). "
     "All effect sizes are small.")

# ---- Table S4: individual factors + BF
rmat = pd.read_csv(TAB / 'phase2_individual_r_matrix_baseline.csv', index_col=0)
qmat = pd.read_csv(TAB / 'phase2_individual_q_matrix_baseline.csv', index_col=0)
bf = pd.read_csv(TAB / 'supp_bayes_factors_individual.csv')
bfs = bf[bf.network == 'SCAN'].set_index('predictor')
labels = {
 'ELA_physical_trauma': 'Physical trauma', 'ELA_caregiver_substance_sep': 'Caregiver substance use / separation',
 'ELA_family_conflict_youth': 'Family conflict', 'ELA_family_aggression': 'Family aggression',
 'ELA_ses_neighborhood': 'SES / neighborhood', 'ELA_caregiver_psych': 'Caregiver psychopathology',
 'ELA_family_anger': 'Family anger', 'ELA_primary_caregiver_support': 'Primary caregiver support',
 'ELA_secondary_caregiver_support': 'Secondary caregiver support', 'ELA_caregiver_supervision': 'Caregiver supervision',
}
# align family anger to the adversity direction (matches Fig 1e / the threat composite)
if 'ELA_family_anger' in rmat.index:
    rmat.loc['ELA_family_anger', 'SCAN'] = -rmat.loc['ELA_family_anger', 'SCAN']
order = rmat['SCAN'].abs().sort_values(ascending=False).index
para('Table S4. Associations of the ten adversity factors with SCAN size, with Bayes factors',
     bold=True, size=11, space_after=4)
rowsS4 = []
for fac in order:
    r = rmat.loc[fac, 'SCAN']; q = qmat.loc[fac, 'SCAN']
    b = bfs.loc[fac]
    rowsS4.append([labels[fac], f"{r:+.3f}", fmt_q(q), fmt_bf(b['BF10'], b['BF01']), b['evidence']])
add_table(['Adversity factor', 'Partial r', 'FDR q', 'Bayes factor', 'Evidence'], rowsS4)
note('Note.', "Partial correlations of each adversity factor with baseline SCAN cortical share, adjusting "
     "for age, sex, head motion, and fixed site. FDR q corrected across the full 10-factor × 15-network "
     "matrix. Bayes factors are BIC approximations from the same models; BF₁₀ > 1 favors an "
     "association, BF₀₁ > 1 favors the null. Evidence labels follow Kass & Raftery (1995).")

# ---- Table S5: all-10 joint
a10 = pd.read_csv(TAB / 'supp_all10factors_SCAN_baseline.csv')
para('Table S5. Joint model of all ten adversity factors predicting SCAN size', bold=True, size=11, space_after=4)
rowsS5 = [[r['label'], f"{r['beta']:+.4f}", f"{r['se']:.4f}", fmt_p(r['p']), fmt_q(r['q_FDR'])]
          for _, r in a10.iterrows()]
add_table(['Adversity factor', 'β', 'SE', 'p', 'FDR q'], rowsS5)
note('Note.', "All ten adversity factors entered simultaneously into a single OLS model predicting "
     "baseline SCAN cortical share (covariates and fixed site as above). Only family aggression remains "
     "independently associated after FDR correction, reflecting shared variance among the threat-type "
     "factors and supporting the composite approach.")

# ---- Table S6: CBCL cross-sectional mediation
med = pd.read_csv(TAB / 'phase6_mediation_SCAN.csv')
cbcl = med[med['outcome'].str.startswith('cbcl')].copy()
n_cbcl = int(cbcl['n'].iloc[0])
para('Table S6. SCAN mediation of threat and year-6 CBCL psychopathology', bold=True, size=11, space_after=4)
rowsS6 = []
for _, r in cbcl.iterrows():
    sig = ' *' if r['mediation_sig'] else ''
    ci = f"[{r['boot_ci_lo']:.3f}, {r['boot_ci_hi']:.3f}]"
    lbl = r['outcome_label'].replace('DSM5', 'DSM-5')
    rowsS6.append([lbl + sig, f"{r['indirect']:.3f}", ci, fmt_p(r['boot_p']),
                   fmt_q(r['q_FDR_indirect'])])
add_table(['CBCL subscale', 'Indirect effect', '95% CI', 'Bootstrap p', 'FDR q'], rowsS6)
note('Note.', f"Indirect effect of threat on each year-6 CBCL subscale through baseline SCAN size "
     f"(n = {n_cbcl:,}), bootstrapped over 5,000 family-clustered resamples with a matched "
     "baseline-subscale covariate. Negative indirect effects would indicate that a larger SCAN is "
     "associated with fewer problems. No subscale reached significance: every 95% confidence interval "
     "includes zero and no indirect effect survives FDR correction across the 14 subscales.")

# ---- Table S7: within-person CBCL
para('Table S7. Within-person SCAN change and CBCL psychopathology', bold=True, size=11, space_after=4)
# Result 2 (ΔSCAN → ΔCBCL), all-waves, all 14 subscales, parsed from the results file.
_wp = WP.read_text()
def _r2(nm):
    m = _re.search(rf"R2 all-waves {_re.escape(nm)}: N=(\d+) \((\d+) subj\) β=([+-][\d.]+)\s+"
                   rf"p=([\d.]+)(?:\s*\*+)?\s+q=([\d.]+)", _wp)
    return m.group(1), m.group(2), m.group(3), float(m.group(4)), float(m.group(5))
rowsS7 = []
_Nobs = _Nsub = None
for _src, _lbl in CBCL_MEDIATION_OUTCOMES.items():
    _Nobs, _Nsub, _b, _p, _q = _r2(_lbl)   # results file labels rows by display name
    rowsS7.append([_lbl.replace('DSM5', 'DSM-5'), _b.replace('-', '−'), fmt_p(_p), fmt_q(_q)])
add_table(['CBCL subscale', 'Within-person β', 'p', 'FDR q'], rowsS7)
note('Note.', f"All-waves within-person models (n = {int(_Nobs):,} observations from {int(_Nsub):,} "
     "children) relating within-person change in SCAN size to within-person change in each of the 14 "
     "CBCL subscales, adjusting for baseline levels, change in head motion, time between waves, baseline "
     "age, and sex, with a subject random intercept and family-cluster-robust standard errors. FDR "
     "corrected across the 14 subscales. No subscale survives correction (all q > .06); the three "
     "nominal associations (p < .05) do not, so within-person SCAN change does not track change in "
     "psychopathology, consistent with the null cross-sectional mediation in Table S6.")

doc.add_page_break()
heading('Supplementary Figures', size=13)
para('Figure S1. Intercorrelations among the ten adversity factors and the three composites',
     bold=True, size=11, space_after=6)
_figp = ROOT / 'outputs/figures/supplement/figS1_ela_correlation.png'
if _figp.exists():
    doc.add_picture(str(_figp), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
note('Note.', "Pearson correlations among the ten z-scored adversity factors (grouped by composite) and the "
     "threat, deprivation, and unpredictability composites, in the analytic sample (n = 4,525). Factors are "
     "shown in composite-aligned orientation, so that higher values on every factor denote greater adversity; "
     "† marks the three factors reverse-coded to achieve this alignment (family anger, primary and secondary "
     "caregiver support). Black lines delineate the composite groupings. Each factor correlates positively with "
     "the composite to which it was assigned (r = 0.41–0.92), and the composites are themselves moderately "
     "intercorrelated, motivating the multivariate specification used in the main text.")

doc.save(str(OUT))
print("Saved:", OUT)
print("Tables:", len(doc.tables), " Paragraphs:", len(doc.paragraphs))
