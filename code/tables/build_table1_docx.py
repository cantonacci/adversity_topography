"""
Build a Word .docx of the sample-characteristics tables, styled to match the
author's example manuscripts (Arial 10pt, borderless banner style, per-block
n/% and M/SD mini-headers, blank-row separators, merged single-value rows).
Numbers are the verified outputs of table1_sample_characteristics.py.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).resolve().parents[2] / 'outputs' / 'tables'
FONT='Arial'; SZ=10

def set_run(run, bold=False, italic=False, size=SZ):
    run.font.name=FONT; run.font.size=Pt(size); run.font.bold=bold; run.font.italic=italic
    # ensure east-asian/complex also Arial
    rpr=run._element.get_or_add_rPr(); rf=rpr.find(qn('w:rFonts'))
    if rf is None:
        rf=OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'),FONT); rf.set(qn('w:hAnsi'),FONT); rf.set(qn('w:cs'),FONT)

def cell_text(cell, text, bold=False, italic=False, align='left', size=SZ):
    cell.text=''
    p=cell.paragraphs[0]
    p.alignment={'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,
                 'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    p.paragraph_format.space_after=Pt(1); p.paragraph_format.space_before=Pt(1)
    run=p.add_run(text); set_run(run,bold,italic,size)

def bottom_border(row, sz='6', color='000000'):
    for cell in row.cells:
        tcPr=cell._tc.get_or_add_tcPr()
        tb=tcPr.find(qn('w:tcBorders'))
        if tb is None:
            tb=OxmlElement('w:tcBorders'); tcPr.append(tb)
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),sz)
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color); tb.append(b)

def caption(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(2)
    run=p.add_run(text); set_run(run,bold=True,size=SZ)

def footnote(text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2)
    run=p.add_run(text); set_run(run,italic=True,size=9)

def main():
    global doc
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    # base style
    nrm = doc.styles['Normal']; nrm.font.name=FONT; nrm.font.size=Pt(SZ)

    # ---------------------------------------------------------------- TABLE 1
    caption("Table 1. Baseline sample characteristics.")
    # rows spec: ('banner',txt) ('kv',label,val) ('hdr',label,c2,c3) ('cat',label,c2,c3) ('blank',)
    T1 = [
     ('banner','Sample Characteristics'),
     ('kv','Children, N','4,525'),
     ('kv','Unique families, N','4,058'),
     ('kv','Children with a sibling in sample, N','923'),
     ('kv','Study sites, N','22'),
     ('blank',),
     ('hdr','Age (years)','M (SD)','Range'),
     ('cat','Baseline','9.98 (0.63)','8.9–11.0'),
     ('kv','Sex','2,171 M, 2,354 F (48.0% / 52.0%)'),
     ('blank',),
     ('hdr','Race/ethnicity','n','%'),
     ('cat','White','2,627','58.1%'),
     ('cat','Hispanic','700','15.5%'),
     ('cat','Other/Multiracial','456','10.1%'),
     ('cat','Black/African American','394','8.7%'),
     ('cat','Asian','64','1.4%'),
     ('cat','Missing','284','6.3%'),
     ('blank',),
     ('hdr','Household income','n','%'),
     ('cat','≥ $100,000','2,327','51.4%'),
     ('cat','$50,000–$99,999','1,116','24.7%'),
     ('cat','< $50,000','798','17.6%'),
     ('cat','Missing','284','6.3%'),
     ('blank',),
     ('hdr','Highest parental education','n','%'),
     ('cat',"Bachelor's degree",'1,384','30.6%'),
     ('cat','Postgraduate degree','1,263','27.9%'),
     ('cat','Some college / associate','1,115','24.6%'),
     ('cat','High-school diploma / GED','336','7.4%'),
     ('cat','< High-school diploma','143','3.2%'),
     ('cat','Missing','284','6.3%'),
     ('blank',),
     ('hdr','Early-life adversity exposures ᵃ','M','SD'),
     ('cat','Threat composite (z)','0.00','1.00'),
     ('cat','Deprivation composite (z)','0.00','1.00'),
     ('cat','Unpredictability composite (z)','0.00','1.00'),
     ('kv','High-threat subgroup (≥ +1 SD), n','579'),
     ('kv','Low-threat subgroup (≤ −1 SD), n','554'),
     ('blank',),
     ('hdr','Imaging and motion QC (baseline rest)','M','SD'),
     ('cat','Mean framewise displacement (mm)','0.118','0.061'),
     ('cat','Retained volumes post-scrubbing','1,189','184'),
     ('cat','Retained rest data (minutes) ᵇ','15.9','2.4'),
     ('hdr','Scanner manufacturer','n','%'),
     ('cat','Siemens','3,012','66.6%'),
     ('cat','GE','1,102','24.4%'),
     ('cat','Philips','411','9.1%'),
     ('blank',),
     ('hdr','SCAN territory (% of cortex)','',''),
     ('cat','Mean (SD)','2.73','(1.12)'),
     ('cat','Median [IQR]','2.51','[1.91–3.31]'),
     ('blank',),
     ('hdr','Cognitive and behavioral outcomes','M','SD'),
     ('cat','NIH-TB Crystallized, baseline (n=4,392)','109.00','18.51'),
     ('cat','NIH-TB Fluid, baseline (n=4,378)','98.92','16.66'),
     ('cat','NIH-TB Crystallized, year 6 (n=2,997)','106.64','15.83'),
     ('cat','NIH-TB Fluid, year 6 (n=2,934)','113.91','19.34'),
     ('cat','CBCL Total Problems (raw)','16.09','16.18'),
     ('cat','CBCL Attention Problems (raw)','2.47','3.14'),
     ('cat','CBCL DSM-5 ADHD (raw)','2.19','2.70'),
     ('cat','CBCL Thought Problems (raw)','1.45','2.00'),
    ]
    tbl = doc.add_table(rows=len(T1), cols=3)
    tbl.style='Normal Table'; tbl.autofit=False
    widths=[Inches(3.4),Inches(1.5),Inches(1.5)]
    for ri,spec in enumerate(T1):
        r=tbl.rows[ri]
        for ci,w in enumerate(widths): r.cells[ci].width=w
        kind=spec[0]
        if kind=='banner':
            merged=r.cells[0].merge(r.cells[1]).merge(r.cells[2])
            cell_text(merged,spec[1],bold=True); bottom_border(r,sz='8')
        elif kind=='blank':
            for c in r.cells: cell_text(c,'')
        elif kind=='kv':
            cell_text(r.cells[0],spec[1])
            merged=r.cells[1].merge(r.cells[2]); cell_text(merged,spec[2],align='left')
        elif kind=='hdr':
            cell_text(r.cells[0],spec[1],bold=True)
            cell_text(r.cells[1],spec[2],italic=True,align='right')
            cell_text(r.cells[2],spec[3],italic=True,align='right')
            bottom_border(r,sz='4',color='808080')
        elif kind=='cat':
            cell_text(r.cells[0],'   '+spec[1])
            cell_text(r.cells[1],spec[2],align='right')
            cell_text(r.cells[2],spec[3],align='right')
    # bottom rule on last row
    bottom_border(tbl.rows[-1],sz='8')
    footnote("Note. N = 4,525 children with usable baseline resting-state data. ᵃ Adversity composites are "
             "standardized (z-scored) across the analytic sample; mean (SD) = 0.00 (1.00) by construction, so the "
             "informative quantities are the ≥/≤ 1 SD subgroup sizes. ᵇ Assuming a repetition time of 0.8 s. "
             "Race/ethnicity, household income, and parental education were missing for the same 284 participants "
             "(6.3%). CBCL scores are raw problem-scale sums; NIH-TB = NIH Toolbox, age-corrected standard scores.")

    # ---------------------------------------------------------------- TABLE 2
    caption("Table 2. Sample retention and age across waves.")
    T2=[('Wave','N','Age, M (SD)','Female, %'),
        ('Baseline (~10 y)','4,525','9.98 (0.63)','52.0'),
        ('Year 2 (~12 y)','4,347','12.03 (0.65)','49.2'),
        ('Year 4 (~14 y)','3,983','14.20 (0.72)','48.4'),
        ('Year 6 (~16 y)','2,687','16.10 (0.66)','47.6')]
    t2=doc.add_table(rows=len(T2),cols=4); t2.style='Normal Table'; t2.autofit=False
    w2=[Inches(2.1),Inches(1.1),Inches(1.6),Inches(1.2)]
    for ri,rowvals in enumerate(T2):
        r=t2.rows[ri]
        for ci,val in enumerate(rowvals):
            r.cells[ci].width=w2[ci]
            cell_text(r.cells[ci],val,bold=(ri==0),align=('left' if ci==0 else 'right'))
        if ri==0: bottom_border(r,sz='8')
    bottom_border(t2.rows[-1],sz='8')
    footnote("Note. N reflects children with usable resting-state topography at each wave.")

    # ---------------------------------------------------------------- SUPP TABLE
    caption("Supplementary Table. Included versus excluded participants (full ABCD baseline cohort, N = 11,870).")
    S=[('','Included (n = 4,525)','Excluded (n = 7,345)','Test'),
       ('Female, %','52.0','45.3','χ²(1) = 50.9, p = 9.6×10⁻¹³, V = 0.07'),
       ('Race/ethnicity, %','','','χ²(4) = 258.2, p = 1.1×10⁻⁵⁴, V = 0.15'),
       ('   White','61.9','48.2',''),
       ('   Hispanic','16.5','21.7',''),
       ('   Black/African American','9.3','17.3',''),
       ('   Other/Multiracial','10.8','10.3',''),
       ('   Asian','1.5','2.5',''),
       ('Household income, %','','','χ²(2) = 251.9, p = 2.0×10⁻⁵⁵, V = 0.16'),
       ('   ≥ $100,000','54.9','41.5',''),
       ('   $50,000–$99,999','26.3','26.9',''),
       ('   < $50,000','18.8','31.5',''),
       ('Parental education, %','','','χ²(4) = 199.2, p = 5.6×10⁻⁴², V = 0.14'),
       ("   Bachelor's degree",'32.6','26.6',''),
       ('   Postgraduate degree','29.8','23.6',''),
       ('   Some college / associate','26.3','30.7',''),
       ('   High-school diploma / GED','7.9','11.6',''),
       ('   < High-school diploma','3.4','7.6','')]
    ts=doc.add_table(rows=len(S),cols=4); ts.style='Normal Table'; ts.autofit=False
    ws=[Inches(2.5),Inches(1.5),Inches(1.5),Inches(2.4)]
    for ri,rowvals in enumerate(S):
        r=ts.rows[ri]
        for ci,val in enumerate(rowvals):
            r.cells[ci].width=ws[ci]
            ishdr=(ri==0)
            cell_text(r.cells[ci],val,bold=ishdr,align=('left' if ci in (0,3) else 'right'))
        if ri==0: bottom_border(r,sz='8')
    bottom_border(ts.rows[-1],sz='8')
    footnote("Note. Percentages are of non-missing values within each column. The included sample is modestly more "
             "White, higher-income, and higher-educated than excluded participants; all effect sizes are small "
             "(bias-corrected Cramér's V ≤ 0.16).")

    path=OUT/'Sample_Characteristics_Tables.docx'
    doc.save(path)
    print("Saved ->", path)


if __name__ == '__main__':
    main()
